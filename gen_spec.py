#!/usr/bin/env python3
"""Generate doc/Spec.docx for PSRAM QSPI Controller IP."""

import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMG_DIR = "doc/image"
os.makedirs(IMG_DIR, exist_ok=True)

# ─────────────────────────────────────────────────────────────
# Helper utilities
# ─────────────────────────────────────────────────────────────
def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def shade_cell(cell, hex_color="D9E1F2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(3)
    return h

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph(text)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    p.paragraph_format.space_after = Pt(4)
    return p

def add_image(doc, path, width_in=5.5, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_in))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph(f"[Image not found: {path}]")

# ─────────────────────────────────────────────────────────────
# Diagram 1 — Clock domain overview
# ─────────────────────────────────────────────────────────────
def draw_clock_domains():
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4)
    ax.axis('off')

    boxes = [
        (0.3, 1.0, 2.4, 2.0, '#AED6F1', 'AXI Domain\n(ACLK)'),
        (3.8, 1.0, 2.4, 2.0, '#A9DFBF', 'APB Domain\n(PCLK)'),
        (7.3, 1.0, 2.4, 2.0, '#F9E79F', 'QSPI Domain\n(SCLK)'),
    ]
    for x, y, w, h, color, label in boxes:
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1",
                                       facecolor=color, edgecolor='#555', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, label, ha='center', va='center',
                fontsize=11, fontweight='bold')

    # Arrow AXI <-> SCLK
    ax.annotate('', xy=(7.3, 2.0), xytext=(2.7, 2.0),
                arrowprops=dict(arrowstyle='<->', color='#C0392B', lw=1.8))
    ax.text(5.0, 2.15, '5× Async FIFO\n(toggle-vector CDC)', ha='center', va='bottom',
            fontsize=8.5, color='#C0392B')

    # Arrow APB <-> SCLK
    ax.annotate('', xy=(7.3, 1.5), xytext=(6.2, 1.5),
                arrowprops=dict(arrowstyle='<->', color='#7D3C98', lw=1.8))
    ax.text(6.75, 1.35, '4-phase\nhandshake', ha='center', va='top',
            fontsize=8, color='#7D3C98')

    ax.set_title('Clock Domain Overview', fontsize=13, fontweight='bold', pad=8)
    path = f"{IMG_DIR}/clock_domains.png"
    fig.savefig(path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return path

# ─────────────────────────────────────────────────────────────
# Diagram 2 — Toggle-vector FIFO CDC
# ─────────────────────────────────────────────────────────────
def draw_toggle_vector():
    fig, ax = plt.subplots(figsize=(11, 5.5))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 5.5)
    ax.axis('off')

    # Write side box
    wr = mpatches.FancyBboxPatch((0.2, 2.5), 3.0, 2.5, boxstyle="round,pad=0.15",
                                  facecolor='#D6EAF8', edgecolor='#2874A6', lw=1.5)
    ax.add_patch(wr)
    ax.text(1.7, 4.9, 'Write Side (i_wr_clk)', ha='center', fontsize=9.5, fontweight='bold', color='#1A5276')
    ax.text(1.7, 4.45, 'reg_wval[N-1:0]', ha='center', fontsize=9, family='monospace')
    ax.text(1.7, 4.05, 'shift left, invert MSB', ha='center', fontsize=8, style='italic')
    ax.text(1.7, 3.6, 'reg_wr_cnt (write ptr)', ha='center', fontsize=9, family='monospace')
    ax.text(1.7, 3.15, 'o_write_valid =\nXNOR(reg_wval, rval_sync)[wr_cnt]', ha='center', fontsize=8)

    # Read side box
    rd = mpatches.FancyBboxPatch((7.8, 2.5), 3.0, 2.5, boxstyle="round,pad=0.15",
                                  facecolor='#D5F5E3', edgecolor='#1E8449', lw=1.5)
    ax.add_patch(rd)
    ax.text(9.3, 4.9, 'Read Side (i_rd_clk)', ha='center', fontsize=9.5, fontweight='bold', color='#145A32')
    ax.text(9.3, 4.45, 'reg_rval[N-1:0]', ha='center', fontsize=9, family='monospace')
    ax.text(9.3, 4.05, 'shift left, invert MSB', ha='center', fontsize=8, style='italic')
    ax.text(9.3, 3.6, 'reg_rd_cnt (read ptr)', ha='center', fontsize=9, family='monospace')
    ax.text(9.3, 3.15, 'o_read_valid =\nXOR(reg_rval, wval_sync)[rd_cnt]', ha='center', fontsize=8)

    # Sync blocks
    ax.add_patch(mpatches.FancyBboxPatch((4.0, 3.8), 2.9, 0.8, boxstyle="round,pad=0.1",
                                          facecolor='#FDEBD0', edgecolor='#CA6F1E', lw=1.2))
    ax.text(5.45, 4.2, '2-stage sync\nrval_sync (on wr_clk)', ha='center', fontsize=8.5)

    ax.add_patch(mpatches.FancyBboxPatch((4.0, 2.7), 2.9, 0.8, boxstyle="round,pad=0.1",
                                          facecolor='#FDEBD0', edgecolor='#CA6F1E', lw=1.2))
    ax.text(5.45, 3.1, '2-stage sync\nwval_sync (on rd_clk)', ha='center', fontsize=8.5)

    # Arrows
    ax.annotate('', xy=(4.0, 4.2), xytext=(3.2, 4.2),
                arrowprops=dict(arrowstyle='->', color='#CA6F1E', lw=1.5))
    ax.annotate('', xy=(3.2, 3.1), xytext=(4.0, 3.1),
                arrowprops=dict(arrowstyle='->', color='#CA6F1E', lw=1.5))
    ax.annotate('', xy=(7.8, 4.2), xytext=(6.9, 4.2),
                arrowprops=dict(arrowstyle='->', color='#CA6F1E', lw=1.5))
    ax.annotate('', xy=(6.9, 3.1), xytext=(7.8, 3.1),
                arrowprops=dict(arrowstyle='->', color='#CA6F1E', lw=1.5))

    # Memory array
    ax.add_patch(mpatches.FancyBboxPatch((3.8, 0.5), 3.3, 1.5, boxstyle="round,pad=0.1",
                                          facecolor='#F2F3F4', edgecolor='#808B96', lw=1.5))
    ax.text(5.45, 1.6, 'reg_mem[0..N-1]', ha='center', fontsize=10, fontweight='bold')
    ax.text(5.45, 1.15, 'Indexed by wr_cnt (write)', ha='center', fontsize=8.5)
    ax.text(5.45, 0.75, 'Indexed by rd_cnt (read)', ha='center', fontsize=8.5)

    ax.annotate('', xy=(3.8, 1.25), xytext=(3.2, 3.2),
                arrowprops=dict(arrowstyle='->', color='#555', lw=1.2, connectionstyle='arc3,rad=0.2'))
    ax.annotate('', xy=(7.1, 1.25), xytext=(7.8, 3.2),
                arrowprops=dict(arrowstyle='->', color='#555', lw=1.2, connectionstyle='arc3,rad=-0.2'))

    ax.set_title('Toggle-Vector Async FIFO CDC Scheme', fontsize=13, fontweight='bold', pad=8)
    path = f"{IMG_DIR}/toggle_vector_fifo.png"
    fig.savefig(path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return path

# ─────────────────────────────────────────────────────────────
# Diagram 3 — APB 4-phase handshake timing
# ─────────────────────────────────────────────────────────────
def draw_apb_cdc_timing():
    fig, axes = plt.subplots(7, 1, figsize=(11, 6), sharex=True)
    fig.subplots_adjust(hspace=0.05, left=0.18, right=0.97)

    t = np.arange(0, 20, 0.01)

    def square(t, edges, val_after):
        y = np.zeros_like(t)
        v = 0
        for i, tt in enumerate(t):
            for e, va in zip(edges, val_after):
                if abs(tt - e) < 0.005:
                    v = va
            y[i] = v
        return y

    signals = [
        ('PCLK',     [1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18], [1,0,1,0,1,0,1,0,1,0,1,0,1,0,1,0,1,0]),
        ('PSEL',     [2.5, 12.0], [1, 0]),
        ('PENABLE',  [3.5, 12.0], [1, 0]),
        ('REQ(PCLK)',[3.5, 11.5], [1, 0]),
        ('SCLK',     [0.3,0.7,1.1,1.5,1.9,2.3,2.7,3.1,3.5,3.9,4.3,4.7,5.1,5.5,5.9,6.3,6.7,7.1,
                      7.5,7.9,8.3,8.7,9.1,9.5,9.9,10.3,10.7,11.1,11.5,11.9,12.3,12.7,13.1,13.5,13.9],
                     [1,0]*18),
        ('ACK(SCLK)',[7.0, 11.0], [1, 0]),
        ('PREADY',   [11.5, 12.5], [1, 0]),
    ]

    colors = ['#2C3E50','#2980B9','#27AE60','#8E44AD','#E67E22','#C0392B','#16A085']

    for ax, (name, edges, vals), color in zip(axes, signals, colors):
        y = square(t, edges, vals)
        ax.step(t, y, where='post', color=color, lw=1.8)
        ax.set_ylim(-0.3, 1.5)
        ax.set_yticks([])
        ax.set_ylabel(name, rotation=0, labelpad=5, ha='right', fontsize=8.5,
                      fontfamily='monospace', color=color)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_visible(False)

    axes[-1].set_xlabel('Time →', fontsize=9)
    axes[-1].set_xticks([])

    # annotation arrows
    for ax in axes:
        ax.axvline(x=3.5, color='gray', lw=0.8, linestyle='--', alpha=0.5)
        ax.axvline(x=7.0, color='gray', lw=0.8, linestyle='--', alpha=0.5)
        ax.axvline(x=11.0, color='gray', lw=0.8, linestyle='--', alpha=0.5)
        ax.axvline(x=11.5, color='gray', lw=0.8, linestyle='--', alpha=0.5)

    axes[0].set_title('APB CDC 4-Phase Handshake Timing (PCLK ↔ SCLK)', fontsize=12, fontweight='bold')
    axes[0].text(3.5, 1.3, '① REQ assert', fontsize=7.5, color='#8E44AD')
    axes[0].text(7.0, 1.3, '② ACK assert', fontsize=7.5, color='#C0392B')
    axes[0].text(11.0, 1.3, '③ ACK fall', fontsize=7.5, color='#16A085')

    path = f"{IMG_DIR}/apb_cdc_timing.png"
    fig.savefig(path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return path

# ─────────────────────────────────────────────────────────────
# Diagram 4 — Register map bit fields (ctrl + key regs)
# ─────────────────────────────────────────────────────────────
def draw_reg_bitfield():
    regs = [
        ('ctrl (0x00)',    [(31,3,'RsvD','#EAECEE'), (2,1,'cmd_2bytes','#AED6F1'),
                            (1,1,'auto_data_wd','#A9DFBF'), (0,1,'en','#F9E79F')]),
        ('wd (0x04)',      [(31,1,'wd_req','#F9E79F'), (30,1,'wd_ack\n(RO)','#EAECEE'),
                            (29,6,'wd_addr','#AED6F1'), (23,24,'wd_data','#A9DFBF')]),
        ('read (0x08)',    [(31,1,'read_reqq','#F9E79F'), (30,1,'read_ack\n(RO)','#EAECEE'),
                            (29,14,'RsvD','#EAECEE'), (15,16,'read_cmd','#AED6F1')]),
        ('write (0x0C)',   [(31,1,'write_reqq','#F9E79F'), (30,1,'write_ack\n(RO)','#EAECEE'),
                            (29,14,'RsvD','#EAECEE'), (15,16,'write_cmd','#AED6F1')]),
        ('wr_dummy(0x10)', [(31,32,'wr_dummy_num[31:0]','#AED6F1')]),
        ('rd_dummy(0x14)', [(31,32,'rd_dummy_num[31:0]','#AED6F1')]),
        ('mode_stat(0x18)',[(31,30,'RsvD','#EAECEE'), (1,2,'current[1:0]','#A9DFBF')]),
        ('indep (0x1C)',   [(31,1,'ind_req','#F9E79F'), (30,1,'ind_ack\n(RO)','#EAECEE'),
                            (29,2,'ind_mode','#AED6F1'), (27,12,'RsvD','#EAECEE'),
                            (15,16,'ind_cmd','#A9DFBF')]),
    ]

    fig, axes = plt.subplots(len(regs), 1, figsize=(12, len(regs)*0.85 + 0.5))
    fig.subplots_adjust(hspace=0.6, left=0.14, right=0.99)

    for ax, (rname, fields) in zip(axes, regs):
        ax.set_xlim(0, 32)
        ax.set_ylim(0, 1)
        ax.axis('off')
        ax.set_ylabel(rname, rotation=0, ha='right', va='center', fontsize=8,
                      fontfamily='monospace')
        x = 32
        for msb, width, label, color in fields:
            lsb = msb - width + 1
            rect = mpatches.Rectangle((32-msb-1, 0.05), width, 0.9,
                                       facecolor=color, edgecolor='#555', lw=0.8)
            ax.add_patch(rect)
            ax.text(32-msb-1 + width/2, 0.5, label, ha='center', va='center',
                    fontsize=6.5 if width > 3 else 5.5, wrap=True)
            ax.text(32-msb-1, 0.02, str(msb), ha='left', va='top', fontsize=5.5, color='#555')
            if width > 1:
                ax.text(32-msb-1+width-0.05, 0.02, str(lsb), ha='right', va='top', fontsize=5.5, color='#555')

    axes[0].set_title('CSR Register Bit-Field Map', fontsize=12, fontweight='bold', y=1.15)

    # legend
    from matplotlib.patches import Patch
    legend_elements = [Patch(facecolor='#F9E79F', edgecolor='#555', label='RWI (req)'),
                       Patch(facecolor='#AED6F1', edgecolor='#555', label='RW'),
                       Patch(facecolor='#A9DFBF', edgecolor='#555', label='RW / RO data'),
                       Patch(facecolor='#EAECEE', edgecolor='#555', label='Reserved / RO')]
    axes[-1].legend(handles=legend_elements, loc='lower right', fontsize=7,
                    bbox_to_anchor=(1.0, -0.5), ncol=4)

    path = f"{IMG_DIR}/reg_bitfields.png"
    fig.savefig(path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return path

# ─────────────────────────────────────────────────────────────
# Diagram 5 — AXI to QSPI data-flow
# ─────────────────────────────────────────────────────────────
def draw_axi_dataflow():
    fig, ax = plt.subplots(figsize=(12, 5))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 5)
    ax.axis('off')

    # Boxes
    def box(ax, x, y, w, h, color, title, subs):
        r = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1",
                                     facecolor=color, edgecolor='#444', lw=1.5)
        ax.add_patch(r)
        ax.text(x+w/2, y+h-0.22, title, ha='center', va='top', fontsize=9, fontweight='bold')
        for i, s in enumerate(subs):
            ax.text(x+w/2, y+h-0.55-i*0.32, s, ha='center', va='top', fontsize=7.5, family='monospace')

    box(ax, 0.1, 1.5, 2.2, 3.0, '#D6EAF8', 'AXI Slave\n(ACLK)',
        ['AW channel','W  channel','AR channel','B  channel','R  channel'])
    box(ax, 3.1, 2.4, 1.8, 2.2, '#FDEBD0', 'FIFOs\n(CDC)',
        ['AWFIFO','WFIFO','ARFIFO','BFIFO','RFIFO'])
    box(ax, 5.5, 1.5, 2.0, 3.0, '#D5F5E3', 'AXI-SCLK\nLogic',
        ['arbiter','req_write','req_read','BRESP gen'])
    box(ax, 8.2, 1.5, 1.8, 3.0, '#F9E79F', 'QSPI FSM\n(SCLK)',
        ['S_IDLE','S_CMD','S_ADDR','S_DUMMY','S_WRITE','S_READ'])
    box(ax, 10.4, 2.2, 1.4, 1.6, '#FDEDEC', 'PSRAM\nDevice',
        ['SIO[3:0]','SCLK','CS#'])

    # Arrows with labels
    def arrow(ax, x1, y1, x2, y2, label, color='#555'):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color=color, lw=1.6))
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx, my+0.12, label, ha='center', fontsize=7.5, color=color)

    arrow(ax, 2.3, 3.2, 3.1, 3.2, 'AW/W/AR')
    arrow(ax, 3.1, 2.7, 2.3, 2.7, 'B/R data')
    arrow(ax, 4.9, 3.2, 5.5, 3.2, 'dequeue')
    arrow(ax, 5.5, 2.7, 4.9, 2.7, 'enqueue B/R')
    arrow(ax, 7.5, 3.2, 8.2, 3.2, 'cmd+addr')
    arrow(ax, 8.2, 2.7, 7.5, 2.7, 'rdata')
    arrow(ax, 10.0, 3.0, 10.4, 3.0, 'QSPI bus')

    # CSR arrow
    box(ax, 3.5, 0.2, 3.5, 0.9, '#E8DAEF', 'CSR / APB (PCLK→SCLK)', ['ctrl, cmd, dummy regs'])
    ax.annotate('', xy=(6.5, 1.5), xytext=(6.5, 1.1),
                arrowprops=dict(arrowstyle='->', color='#7D3C98', lw=1.4))
    ax.text(6.7, 1.3, 'config', fontsize=7.5, color='#7D3C98')

    ax.set_title('AXI → QSPI Data-Flow Overview', fontsize=13, fontweight='bold', pad=6)
    path = f"{IMG_DIR}/axi_dataflow.png"
    fig.savefig(path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return path

print("Generating diagrams...")
p_clk   = draw_clock_domains()
p_tvcdc = draw_toggle_vector()
p_apb   = draw_apb_cdc_timing()
p_reg   = draw_reg_bitfield()
p_flow  = draw_axi_dataflow()
print("Diagrams done.")
print(p_clk, p_tvcdc, p_apb, p_reg, p_flow)
