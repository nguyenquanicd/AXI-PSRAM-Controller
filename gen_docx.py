#!/usr/bin/env python3
"""Build doc/Spec.docx for PSRAM QSPI Controller IP."""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMG = "doc/image"
OUT = "doc/Spec.docx"

# ── helpers ──────────────────────────────────────────────────
def shade(cell, hex_color="D9E1F2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def hdr_row(table, *texts, color="2E4057"):
    row = table.rows[0]
    for cell, text in zip(row.cells, texts):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, color)

def add_img(doc, fname, caption, width=5.8):
    path = f"{IMG}/{fname}"
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph(f"[missing: {path}]")

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p

# ── document setup ────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width  = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.54)
sec.top_margin  = sec.bottom_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ── Title page ────────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph("PSRAM QSPI Controller IP")
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.size = Pt(24)
t.runs[0].bold = True

sub = doc.add_paragraph("Hardware Design Specification")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(16)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Version: 1.0    |    Date: 2026-06-23    |    Status: Draft\n")
meta.add_run("Audience: SoC Integration Engineers")
meta.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# SECTION I — Overview
# ═══════════════════════════════════════════════════════════
doc.add_heading("I. Overview", 1)
body(doc,
    "The PSRAM QSPI Controller is a synthesizable RTL IP that bridges an AXI4 full slave "
    "interface to an external Pseudo-SRAM (PSRAM) device using a Quad-SPI (QSPI) or single "
    "SPI bus. It is designed for SoC integration where the PSRAM serves as extended on-chip "
    "memory accessible from a high-performance AXI4 fabric."
)
body(doc,
    "The IP supports full AXI4 burst transfers (FIXED, INCR, WRAP), independent command "
    "injection via an APB slave interface, XIP-style multi-beat reads, and configurable "
    "dummy cycles. Three independent clock domains are supported with full CDC treatment."
)

doc.add_heading("I.1  Key Features", 2)
features = [
    "AXI4 full slave: burst types FIXED / INCR / WRAP, up to 256-beat bursts (AWLEN/ARLEN 8-bit)",
    "APB slave for CSR configuration — fully decoupled from data path",
    "Three independent clock domains: ACLK (AXI), PCLK (APB), SCLK (QSPI bus)",
    "CDC ACLK↔SCLK via 5 async FIFOs using a toggle-vector scheme",
    "CDC PCLK↔SCLK via 4-phase handshake with 2-stage synchronizers",
    "Configurable QSPI / SPI mode; 1-byte or 2-byte command width",
    "Automatic data-width selection (auto_data_wd)",
    "Configurable dummy cycles per operation (32-bit register, FSM uses [3:0])",
    "Round-robin write/read arbiter to prevent starvation",
    "BRESP = SLVERR on AW/W burst-length mismatch",
]
for f in features:
    bullet(doc, f)

doc.add_heading("I.2  Parameters", 2)
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hdr_row(tbl, "Parameter", "Default", "Description")
params = [
    ("PARA_AXI_DATA_WD", "32", "AXI data bus width in bits"),
    ("PARA_AXI_ADDR_WD", "24", "AXI address width in bits"),
    ("PARA_AXI_ID_WD",   "4",  "AXI ID field width"),
    ("PARA_AXI_LEN_WD",  "8",  "AXI burst length field width (AWLEN/ARLEN)"),
    ("PARA_AXI_FIFO_DEPTH","8","Async FIFO depth (must be power of 2)"),
]
for name, default, desc in params:
    row = tbl.add_row()
    row.cells[0].text = name
    row.cells[1].text = default
    row.cells[2].text = desc
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)

doc.add_paragraph()
add_img(doc, "psram_overview.png", "Figure 1 — Top-level block diagram")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# SECTION IV — Clock Domain Crossing (CDC)
# ═══════════════════════════════════════════════════════════
doc.add_heading("IV. Clock Domain Crossing (CDC)", 1)

doc.add_heading("IV.1  ACLK ↔ SCLK: Toggle-Vector Async FIFO", 2)
body(doc,
    "Five async FIFOs decouple the AXI domain from the QSPI domain: AWFIFO, WFIFO (write "
    "path), ARFIFO (read path), BFIFO (write response), and RFIFO (read data). Each FIFO "
    "uses a toggle-vector occupancy scheme — not Gray-encoded pointers."
)
add_img(doc, "toggle_vector_fifo.png", "Figure 3 — Toggle-vector async FIFO CDC structure")
doc.add_paragraph()

body(doc, "Write-side occupancy detection:")
bullet(doc, "reg_wval is a PARA_DEPTH-bit shift register. Each write shifts left and inverts the MSB: reg_wval <= {reg_wval[N-2:0], ~reg_wval[N-1]}.")
bullet(doc, "rval_sync is reg_rval synchronized into the write clock domain via a PARA_DEPTH-wide 2-stage synchronizer (m_vlsi_multi_synch).")
bullet(doc, "wr_comb_val = reg_wval XNOR rval_sync. o_write_valid = wr_comb_val[reg_wr_cnt].")
bullet(doc, "A slot is writable when the write-side and read-side vectors agree at that index position.")

body(doc, "Read-side occupancy detection:")
bullet(doc, "reg_rval is a PARA_DEPTH-bit shift register. Each read shifts left and inverts the MSB.")
bullet(doc, "wval_sync is reg_wval synchronized into the read clock domain.")
bullet(doc, "rd_comb_val = reg_rval XOR wval_sync. o_read_valid = rd_comb_val[reg_rd_cnt].")
bullet(doc, "A slot is readable when the read-side and write-side vectors differ at that index position.")

doc.add_heading("IV.2  PCLK ↔ SCLK: 4-Phase Handshake", 2)
body(doc,
    "APB register accesses cross from PCLK to SCLK via a 4-phase REQ/ACK handshake. "
    "Each direction uses a pair of 2-stage synchronizers (m_vlsi_synch, PARA_LEVELS=2)."
)
add_img(doc, "apb_cdc_timing.png", "Figure 4 — APB CDC 4-phase handshake timing diagram")
doc.add_paragraph()

body(doc, "Handshake sequence for a write or read access:")
bullet(doc, "① PCLK side asserts REQ on detection of APB setup phase (PSEL & ~PENABLE, no SLVERR).")
bullet(doc, "② SCLK side detects rising edge of synchronized REQ, performs the register access, then asserts ACK.")
bullet(doc, "③ PCLK side detects rising edge of synchronized ACK; PREADY is registered one PCLK cycle after the ACK falling edge.")
bullet(doc, "④ PCLK side de-asserts REQ; SCLK side de-asserts ACK — handshake complete.")

body(doc,
    "PREADY is asserted for one PCLK cycle on the falling edge of the synchronized ACK signal, "
    "or immediately when SLVERR is detected in the setup phase. The APB master must not "
    "issue a back-to-back transaction until PREADY is observed."
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
doc.add_heading("III. Clocks and Resets", 1)
body(doc,
    "The IP operates across three fully independent clock domains. Each domain has its own "
    "clock and active-low synchronous reset. There is no internal clock gating or division; "
    "all clock sources must be supplied by the SoC."
)
add_img(doc, "clock_domains.png", "Figure 2 — Clock domain partitioning and CDC paths")
doc.add_paragraph()

# Clock summary table
tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
hdr_row(tbl, "Clock", "Signal", "Reset", "Description")
clk_rows = [
    ("ACLK",  "i_axi_clk",  "i_axi_rstn",  "Clocks the AXI4 slave interface and async-FIFO write/read ports on the AXI side"),
    ("PCLK",  "i_apb_clk",  "i_apb_rstn",  "Clocks the APB slave interface and the 4-phase handshake bus-side logic"),
    ("SCLK",  "i_qspi_clk", "i_qspi_rstn", "Clocks the QSPI FSM, async-FIFO SCLK-side ports, and CSR register storage"),
]
for clk, sig, rst, desc in clk_rows:
    r = tbl.add_row()
    r.cells[0].text = clk
    r.cells[1].text = sig
    r.cells[2].text = rst
    r.cells[3].text = desc
    for c in r.cells:
        c.paragraphs[0].runs[0].font.size = Pt(9)
    r.cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
    r.cells[2].paragraphs[0].runs[0].font.name = 'Courier New'
doc.add_paragraph()

body(doc,
    "The QSPI clock is wired directly to the pad: o_qspi_sclk = i_qspi_clk. "
    "No gating, division, or phase shift is applied. The SoC must ensure i_qspi_clk is "
    "disabled when the PSRAM is not selected to avoid spurious clocking."
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
doc.add_heading("II. Port Description", 1)

def port_table(doc, title, ports):
    doc.add_heading(title, 2)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    hdr_row(tbl, "Signal", "Dir", "Width", "Description")
    for sig, direction, width, desc in ports:
        row = tbl.add_row()
        row.cells[0].text = sig
        row.cells[1].text = direction
        row.cells[2].text = width
        row.cells[3].text = desc
        row.cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

port_table(doc, "II.1  Clock and Reset", [
    ("i_axi_clk",   "IN",  "1", "AXI domain clock (ACLK)"),
    ("i_axi_rstn",  "IN",  "1", "AXI domain active-low synchronous reset"),
    ("i_apb_clk",   "IN",  "1", "APB domain clock (PCLK)"),
    ("i_apb_rstn",  "IN",  "1", "APB domain active-low synchronous reset"),
    ("i_qspi_clk",  "IN",  "1", "QSPI clock source; wired directly to o_qspi_sclk (no gating)"),
    ("i_qspi_rstn", "IN",  "1", "QSPI domain active-low synchronous reset"),
])

port_table(doc, "II.2  AXI4 Write Address Channel (AW)", [
    ("i_awvalid", "IN",  "1",                "Write address valid"),
    ("o_awready", "OUT", "1",                "Write address ready"),
    ("i_awid",    "IN",  "PARA_AXI_ID_WD",   "Write address ID"),
    ("i_awaddr",  "IN",  "PARA_AXI_ADDR_WD", "Write start address"),
    ("i_awlen",   "IN",  "PARA_AXI_LEN_WD",  "Burst length – 1 (0 = 1 beat, 255 = 256 beats)"),
    ("i_awsize",  "IN",  "3",                "Transfer size: 2^awsize bytes per beat"),
    ("i_awburst", "IN",  "2",                "Burst type: 00=FIXED, 01=INCR, 10=WRAP"),
])

port_table(doc, "II.3  AXI4 Write Data Channel (W)", [
    ("i_wvalid", "IN",  "1",                    "Write data valid"),
    ("o_wready", "OUT", "1",                    "Write data ready"),
    ("i_wdata",  "IN",  "PARA_AXI_DATA_WD",     "Write data"),
    ("i_wstrb",  "IN",  "PARA_AXI_DATA_WD/8",   "Byte-enable strobes"),
    ("i_wlast",  "IN",  "1",                    "Asserted on final beat of write burst"),
])

port_table(doc, "II.4  AXI4 Write Response Channel (B)", [
    ("o_bvalid", "OUT", "1",               "Write response valid"),
    ("i_bready", "IN",  "1",               "Write response ready"),
    ("o_bid",    "OUT", "PARA_AXI_ID_WD",  "Write response ID (echoes AWID)"),
    ("o_bresp",  "OUT", "2",               "Response code: 2'b00=OKAY, 2'b10=SLVERR"),
])

port_table(doc, "II.5  AXI4 Read Address Channel (AR)", [
    ("i_arvalid", "IN",  "1",                "Read address valid"),
    ("o_arready", "OUT", "1",                "Read address ready"),
    ("i_arid",    "IN",  "PARA_AXI_ID_WD",   "Read address ID"),
    ("i_araddr",  "IN",  "PARA_AXI_ADDR_WD", "Read start address"),
    ("i_arlen",   "IN",  "PARA_AXI_LEN_WD",  "Burst length – 1"),
    ("i_arsize",  "IN",  "3",                "Transfer size"),
    ("i_arburst", "IN",  "2",                "Burst type"),
])

port_table(doc, "II.6  AXI4 Read Data Channel (R)", [
    ("o_rvalid", "OUT", "1",               "Read data valid"),
    ("i_rready", "IN",  "1",               "Read data ready"),
    ("o_rid",    "OUT", "PARA_AXI_ID_WD",  "Read data ID (echoes ARID)"),
    ("o_rdata",  "OUT", "PARA_AXI_DATA_WD","Read data"),
    ("o_rresp",  "OUT", "2",               "Read response (always 2'b00=OKAY)"),
    ("o_rlast",  "OUT", "1",               "Asserted on final beat of read burst"),
])

port_table(doc, "II.7  APB Slave Interface", [
    ("i_psel",        "IN",  "1",  "APB select"),
    ("i_penable",     "IN",  "1",  "APB enable (access phase strobe)"),
    ("i_pwrite",      "IN",  "1",  "1 = write transaction, 0 = read transaction"),
    ("i_paddr",       "IN",  "16", "APB address (byte-addressed, must be 4-byte aligned)"),
    ("i_pwdata",      "IN",  "32", "APB write data"),
    ("i_pstrb",       "IN",  "4",  "Write byte-enable strobes; all 4 must be asserted"),
    ("o_prdata",      "OUT", "32", "APB read data"),
    ("o_pready",      "OUT", "1",  "APB ready; de-asserts to stall during cross-domain transfer"),
    ("o_pslverr",     "OUT", "1",  "APB slave error (protection violation / address error)"),
    ("i_pprot",       "IN",  "3",  "APB protection type (pprot[1] checked when protect_en=1)"),
    ("i_protect_en",  "IN",  "1",  "Enable protection check on pprot[1]"),
    ("i_slverr_en",   "IN",  "1",  "Enable slave-error generation"),
])

port_table(doc, "II.8  QSPI Pad Interface", [
    ("o_qspi_sclk", "OUT",   "1", "QSPI serial clock — combinational passthrough of i_qspi_clk"),
    ("o_qspi_csn",  "OUT",   "1", "QSPI chip-select, active low"),
    ("io_qspi_sio", "INOUT", "4", "Quad-SPI data lines: SIO[0]=MOSI, SIO[1]=MISO in SPI mode"),
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# SECTION V — CSR Register Map
# ═══════════════════════════════════════════════════════════
doc.add_heading("V. CSR Register Map", 1)
body(doc,
    "All CSR registers are 32 bits wide and accessed via the APB slave interface. "
    "The base address is SoC-dependent. All offsets are byte-addressed and 4-byte aligned. "
    "Register flip-flops are clocked by i_reg_clk (SCLK domain); the APB bus uses i_bus_clk (PCLK domain)."
)
add_img(doc, "reg_bitfields.png", "Figure 5 — CSR bit-field map for all 8 registers")
doc.add_paragraph()

reg_tbl = doc.add_table(rows=1, cols=4)
reg_tbl.style = 'Table Grid'
hdr_row(reg_tbl, "Offset", "Name", "Access", "Description")
for offset, name, access, desc in [
    ("0x00", "ctrl",        "RW",  "Control: en, auto_data_wd, cmd_2bytes"),
    ("0x04", "wd",          "RWI", "Write-data command: req, ack(RO), addr[5:0], data[23:0]"),
    ("0x08", "read",        "RWI", "Read command override: req, ack(RO), read_cmd[15:0]"),
    ("0x0C", "write",       "RWI", "Write command override: req, ack(RO), write_cmd[15:0]"),
    ("0x10", "wr_dummy",    "RW",  "Write dummy cycles [31:0]; FSM uses [3:0] only"),
    ("0x14", "rd_dummy",    "RW",  "Read dummy cycles [31:0]; FSM uses [3:0] only"),
    ("0x18", "mode_status", "RO",  "Current QSPI mode [1:0] — combinational from i_mode_status_current"),
    ("0x1C", "independent", "RWI", "Independent cmd injection: req, ack(RO), mode[1:0], cmd[15:0]"),
]:
    r = reg_tbl.add_row()
    for i, v in enumerate([offset, name, access, desc]):
        r.cells[i].text = v
        r.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    r.cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
    r.cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
doc.add_paragraph()
